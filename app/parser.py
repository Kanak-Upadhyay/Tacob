from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from pdfminer.high_level import extract_text

from app.skills import ALL_SKILLS, ROLE_FAMILIES, TITLE_HINTS


EMAIL_RE = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.I)
PHONE_RE = re.compile(r"(?:\+?\d{1,3}[\s-]?)?(?:\(?\d{2,4}\)?[\s-]?)?\d{3,5}[\s-]?\d{4,6}")
LINKEDIN_RE = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9\-_/]+", re.I)
GITHUB_RE = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9\-]+", re.I)
URL_RE = re.compile(r"https?://[^\s)]+", re.I)
YEAR_RE = re.compile(r"(20\d{2}|19\d{2})")
METRIC_RE = re.compile(
    r"(\d+(?:\.\d+)?\s*%|₹\s*\d[\d,]*(?:\s*(?:lakh|lac|cr|crore))?|\b\d+\s*(?:x|times)\b)",
    re.I,
)

SECTION_ALIASES = {
    "summary": ["summary", "professional summary", "profile", "about", "objective", "career objective"],
    "skills": ["skills", "technical skills", "core competencies", "key skills", "expertise", "tech stack"],
    "experience": [
        "experience", "work experience", "professional experience", "employment",
        "work history", "career history",
    ],
    "education": ["education", "academic", "academics", "qualifications"],
    "projects": ["projects", "personal projects", "academic projects"],
    "certifications": ["certifications", "certificates", "licenses"],
    "achievements": ["achievements", "accomplishments", "awards"],
}


@dataclass
class ParsedCV:
    filename: str
    text: str
    word_count: int
    name: str
    emails: list[str]
    phones: list[str]
    linkedin: str | None
    github: str | None
    location: str | None
    headline: str | None
    skills: list[str]
    sections: dict[str, str]
    titles: list[str]
    years_experience: float | None
    education: list[str]
    metrics_found: list[str]
    likely_scanned: bool
    role_family: str | None
    sample: bool = False
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        return {
            "filename": self.filename,
            "text": self.text,
            "word_count": self.word_count,
            "name": self.name,
            "emails": self.emails,
            "phones": self.phones,
            "linkedin": self.linkedin,
            "github": self.github,
            "location": self.location,
            "headline": self.headline,
            "skills": self.skills,
            "sections": {k: v[:4000] for k, v in self.sections.items()},
            "titles": self.titles,
            "years_experience": self.years_experience,
            "education": self.education,
            "metrics_found": self.metrics_found[:12],
            "likely_scanned": self.likely_scanned,
            "role_family": self.role_family,
            "sample": self.sample,
            "warnings": self.warnings,
        }


def extract_text_from_upload(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return extract_text(io.BytesIO(data)) or ""
    if suffix in {".docx"}:
        doc = Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        return "\n".join(parts)
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="ignore")
    raise ValueError("Please upload a PDF, DOCX, or TXT resume.")


def parse_cv_text(text: str, filename: str = "resume.txt", sample: bool = False) -> ParsedCV:
    cleaned = _normalize(text)
    warnings: list[str] = []
    if not cleaned.strip():
        raise ValueError("Could not read any text from this file.")

    words = re.findall(r"[A-Za-z0-9']+", cleaned)
    likely_scanned = len(words) < 80
    if likely_scanned:
        warnings.append("Very little text was extracted. This may be a scanned/image PDF, which most ATS systems reject.")

    emails = _unique(EMAIL_RE.findall(cleaned))
    phones = [p.strip() for p in PHONE_RE.findall(cleaned) if len(re.sub(r"\D", "", p)) >= 10][:3]
    linkedin_match = LINKEDIN_RE.search(cleaned)
    github_match = GITHUB_RE.search(cleaned)
    sections = _split_sections(cleaned)
    skills = _extract_skills(cleaned, sections)
    titles = _extract_titles(cleaned, sections.get("experience", ""))
    education = _extract_education(sections.get("education", cleaned[:2000]))
    metrics = METRIC_RE.findall(cleaned)
    name = _guess_name(cleaned, emails)
    headline = _guess_headline(cleaned, titles)
    location = _guess_location(cleaned)
    years = _guess_years(cleaned)
    family = _guess_family(skills, titles, cleaned)

    if not emails:
        warnings.append("No email address found.")
    if not phones:
        warnings.append("No phone number found.")
    if not linkedin_match:
        warnings.append("No LinkedIn URL found.")

    return ParsedCV(
        filename=filename,
        text=cleaned,
        word_count=len(words),
        name=name,
        emails=emails,
        phones=phones,
        linkedin=linkedin_match.group(0) if linkedin_match else None,
        github=github_match.group(0) if github_match else None,
        location=location,
        headline=headline,
        skills=skills,
        sections=sections,
        titles=titles,
        years_experience=years,
        education=education,
        metrics_found=[m if isinstance(m, str) else m[0] for m in metrics],
        likely_scanned=likely_scanned,
        role_family=family,
        sample=sample,
        warnings=warnings,
    )


def _normalize(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _unique(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for item in items:
        key = item.lower()
        if key not in seen:
            seen.add(key)
            out.append(item)
    return out


GENERIC_SKIP = {
    "inventory", "operations", "leadership", "communication", "training",
    "presentation", "sales",
}


def _extract_skills(text: str, sections: dict[str, str] | None = None) -> list[str]:
    lower = text.lower()
    skills_blob = (sections or {}).get("skills", "").lower()
    found: list[str] = []
    for skill in ALL_SKILLS:
        pattern = r"(?<![A-Za-z0-9])" + re.escape(skill) + r"(?![A-Za-z0-9])"
        if not re.search(pattern, lower):
            continue
        if skill in GENERIC_SKIP and skills_blob and not re.search(pattern, skills_blob):
            continue
        found.append(skill)
    # Keep original-ish display: title-case unless it contains punctuation
    display = []
    for skill in found:
        if any(ch in skill for ch in ".+#"):
            display.append(skill)
        else:
            display.append(skill.title() if " " in skill else skill)
    return _unique(display)[:40]


def _split_sections(text: str) -> dict[str, str]:
    lines = text.splitlines()
    headings: list[tuple[int, str]] = []
    for i, line in enumerate(lines):
        key = _match_heading(line)
        if key:
            headings.append((i, key))
    if not headings:
        return {}
    sections: dict[str, str] = {}
    for idx, (start, key) in enumerate(headings):
        end = headings[idx + 1][0] if idx + 1 < len(headings) else len(lines)
        body = "\n".join(lines[start + 1 : end]).strip()
        sections[key] = (sections.get(key, "") + "\n" + body).strip()
    return sections


def _match_heading(line: str) -> str | None:
    stripped = re.sub(r"[^A-Za-z /&]", "", line).strip().lower()
    if not stripped or len(stripped) > 42:
        return None
    for key, aliases in SECTION_ALIASES.items():
        if stripped in aliases:
            return key
    return None


def _extract_titles(full_text: str, experience: str) -> list[str]:
    blob = (experience or full_text).lower()
    found: list[str] = []
    for hint in TITLE_HINTS:
        if hint in blob:
            found.append(hint.title() if hint != "ui/ux" else "UI/UX Designer")
    # Also pick lines that look like job titles near dates
    for line in (experience or full_text).splitlines():
        if YEAR_RE.search(line) and 8 < len(line) < 80:
            for hint in TITLE_HINTS:
                if hint in line.lower() and hint.title() not in found:
                    found.append(hint.title())
    return _unique(found)[:8]


def _extract_education(text: str) -> list[str]:
    degrees = []
    patterns = [
        r"b\.?\s*tech[^,\n]{0,40}",
        r"m\.?\s*tech[^,\n]{0,40}",
        r"b\.?\s*e\.?[^,\n]{0,40}",
        r"m\.?\s*c\.?\s*a\.?[^,\n]{0,30}",
        r"b\.?\s*c\.?\s*a\.?[^,\n]{0,30}",
        r"mba[^,\n]{0,40}",
        r"b\.?\s*com[^,\n]{0,40}",
        r"m\.?\s*com[^,\n]{0,40}",
        r"b\.?\s*sc[^,\n]{0,40}",
        r"bachelor[^,\n]{0,50}",
        r"master[^,\n]{0,50}",
        r"ph\.?\s*d[^,\n]{0,40}",
    ]
    lower = text
    for pat in patterns:
        for match in re.finditer(pat, lower, re.I):
            degrees.append(re.sub(r"\s+", " ", match.group(0)).strip())
    return _unique(degrees)[:6]


def _guess_name(text: str, emails: list[str]) -> str:
    for line in text.splitlines()[:8]:
        candidate = line.strip()
        if not candidate or EMAIL_RE.search(candidate) or PHONE_RE.search(candidate):
            continue
        if URL_RE.search(candidate) or "linkedin" in candidate.lower():
            continue
        words = candidate.split()
        if 1 < len(words) <= 4 and all(w[0].isalpha() for w in words if w):
            if not any(k in candidate.lower() for k in ("resume", "curriculum", "cv")):
                return candidate.title() if candidate.isupper() else candidate
    if emails:
        local = emails[0].split("@")[0]
        return re.sub(r"[._-]+", " ", local).title()
    return "Candidate"


def _guess_headline(text: str, titles: list[str]) -> str | None:
    if titles:
        return titles[0]
    for line in text.splitlines()[1:10]:
        low = line.strip().lower()
        if 12 < len(low) < 70 and not EMAIL_RE.search(line) and "skill" not in low:
            if any(h in low for h in TITLE_HINTS):
                return line.strip()
    return None


def _guess_location(text: str) -> str | None:
    cities = [
        "Bengaluru", "Bangalore", "Hyderabad", "Pune", "Mumbai", "Delhi",
        "Noida", "Gurgaon", "Gurugram", "Chennai", "Kolkata", "Ahmedabad",
        "Jaipur", "Kochi", "Indore", "Chandigarh", "Remote",
    ]
    lower = text.lower()
    for city in cities:
        if city.lower() in lower:
            return "Bengaluru" if city == "Bangalore" else city
    return None


def _guess_years(text: str) -> float | None:
    match = re.search(r"(\d+(?:\.\d+)?)\+?\s*(?:years?|yrs)\s+(?:of\s+)?(?:experience|exp)", text, re.I)
    if match:
        return float(match.group(1))
    years = [int(y) for y in YEAR_RE.findall(text)]
    years = [y for y in years if 1995 <= y <= 2026]
    if len(years) >= 2:
        span = max(years) - min(years)
        if 0 < span <= 40:
            return float(span)
    return None


def _guess_family(skills: list[str], titles: list[str], text: str) -> str | None:
    skill_set = {s.lower() for s in skills}
    blob = text.lower() + " " + " ".join(t.lower() for t in titles)
    best_key = None
    best_score = 0
    for family in ROLE_FAMILIES:
        score = sum(1 for k in family["keywords"] if k in skill_set or k in blob)
        for title in family["titles"]:
            if title.lower() in blob:
                score += 3
        if score > best_score:
            best_score = score
            best_key = family["key"]
    return best_key if best_score >= 2 else None
